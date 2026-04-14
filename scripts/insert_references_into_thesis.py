from __future__ import annotations

from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "output" / "thesis_cn_visual_final_fitted.docx"
OUT = ROOT / "output" / "thesis_cn_visual_final_referenced.docx"


NEW_REFERENCES = [
    "[1] Fortunato S. Community detection in graphs[J]. Physics Reports, 2010, 486(3-5): 75-174.",
    "[2] Blondel V D, Guillaume J L, Lambiotte R, et al. Fast unfolding of communities in large networks[J]. Journal of Statistical Mechanics: Theory and Experiment, 2008(10): P10008.",
    "[3] Lancichinetti A, Fortunato S, Radicchi F. Benchmark graphs for testing community detection algorithms[J]. Physical Review E, 2008, 78(4): 046110.",
    "[4] Mucha P J, Richardson T, Macon K, et al. Community structure in time-dependent, multiscale, and multiplex networks[J]. Science, 2010, 328(5980): 876-878.",
    "[5] Li J, Lai S, Shuai Z, et al. A comprehensive review of community detection in graphs[J]. Neurocomputing, 2024, 600: 128169.",
    "[6] Kivelä M, Arenas A, Barthélemy M, et al. Multilayer networks[J]. Journal of Complex Networks, 2014, 2(3): 203-271.",
    "[7] Zhu T, Li G, Zhou W, et al. Differentially private data publishing and analysis: A survey[J]. IEEE Transactions on Knowledge and Data Engineering, 2017, 29(8): 1619-1638.",
    "[8] Dwork C, Roth A. The Algorithmic Foundations of Differential Privacy[J]. Foundations and Trends in Theoretical Computer Science, 2014, 9(3-4): 211-407.",
    "[9] Paillier P. Public-key cryptosystems based on composite degree residuosity classes[C]// Advances in Cryptology—EUROCRYPT'99. Berlin: Springer, 1999: 223-238.",
    "[10] Ma L, Huang X, Li J, et al. Privacy-preserving global structural balance computation in signed networks[J]. IEEE Transactions on Computational Social Systems, 2019, 7(1): 164-177.",
    "[11] Shao Z, Ma L, Lin Q, et al. PMCDM: Privacy-preserving multiresolution community detection in multiplex networks[J]. Knowledge-Based Systems, 2022, 244: 108542.",
    "[12] Fu N, Ni W, Hou L, et al. Community detection in decentralized social networks with local differential privacy[J]. Information Sciences, 2024, 661: 120164.",
    "[13] Guo K, Chen D, Huang Q, et al. Privacy-preserving multi-label propagation based on federated learning[J]. IEEE Transactions on Network Science and Engineering, 2023, 11(1): 886-899.",
    "[14] Ferrag M A, Maglaras L, Ahmim A. Privacy-preserving schemes for ad hoc social networks: A survey[J]. IEEE Communications Surveys & Tutorials, 2017, 19(4): 3015-3045.",
    "[15] Oughtred R, Rust J, Chang C, et al. The BioGRID interaction database: 2019 update[J]. Nucleic Acids Research, 2019, 47(D1): D529-D541.",
    "[16] Magnani M, Micenková B, Rossi L. Combinatorial analysis of multiple networks[EB/OL]. arXiv, 2013: arXiv:1303.4986.",
    "[17] Bródka P. A Method for Group Extraction and Analysis in Multilayer Social Networks[EB/OL]. arXiv, 2016: arXiv:1612.02377.",
]


def replace_para_text(para, text: str) -> None:
    para.clear()
    para.add_run(text)


def main() -> None:
    doc = Document(SRC)

    for para in doc.paragraphs:
        text = para.text.strip()

        if text.startswith("国外关于社区检测的研究起步较早。Fortunato"):
            replace_para_text(
                para,
                "国外关于社区检测的研究起步较早。Fortunato 对社区检测问题进行了系统综述，指出社区结构能够帮助理解复杂系统的组织方式，并总结了模块度优化、谱聚类、标签传播等多类方法[1]。Blondel 等人提出的 Louvain 算法由于具有较好的效率和可扩展性，成为大规模网络社区检测中的经典方法[2]。在测试方法方面，Lancichinetti、Fortunato 与 Radicchi 提出的 LFR 基准网络更接近真实网络的度分布和社区规模分布，现已成为评价社区检测算法的重要基准[3]。近年来，Li 等人又从综述角度对图社区检测方法进行了系统归纳，进一步说明了该研究方向仍具有持续发展价值[5]。",
            )
        elif text.startswith("国内关于复杂网络社区发现的研究更多集中在算法改进"):
            replace_para_text(
                para,
                "国内关于复杂网络社区发现的研究更多集中在算法改进、指标优化以及特定应用场景建模等方面。随着社交网络分析、生物网络分析和多源异构数据分析需求的增加，围绕模块度缺陷修正、重叠社区识别以及多层网络划分的研究逐渐增多，但与国外相比，面向隐私保护与系统实现一体化的研究成果仍相对有限[5]。",
            )
        elif text.startswith("随着多源关系网络的普及，研究者开始关注多层网络和多路复用网络中的社区结构。"):
            replace_para_text(
                para,
                "随着多源关系网络的普及，研究者开始关注多层网络和多路复用网络中的社区结构。Mucha 等人较早从时变、多尺度和多路复用网络的统一建模角度讨论了社区结构问题[4]；Kivelä 等人进一步总结了多层网络的表示方法与分析框架，指出传统单层图模型难以完整刻画多种关系并存的复杂系统[6]。因此，在多层网络场景下，如何综合各层结构信息进行社区划分，成为社区检测研究的重要延伸方向。",
            )
        elif text.startswith("国外关于隐私保护数据分析的研究体系较为成熟。"):
            replace_para_text(
                para,
                "国外关于隐私保护数据分析的研究体系较为成熟。Dwork 与 Roth 系统阐述了差分隐私的定义、性质及其机制设计方法，为隐私保护数据分析奠定了理论基础[8]；Zhu 等人则从数据发布与分析角度总结了差分隐私的主要研究脉络[7]。Paillier 提出的加法同态公钥密码体制，使得在密文状态下进行加法和数乘成为可能，为安全统计与隐私计算提供了工具[9]。这些研究为图数据隐私保护提供了重要理论支撑。",
            )
        elif text.startswith("国内在隐私保护图分析方面也已有探索"):
            replace_para_text(
                para,
                "国内外在隐私保护图分析方面也已有探索，研究内容涉及社交网络匿名化、图发布保护、联邦图学习和差分隐私机制改进等方向。Ferrag 等人对社会网络隐私保护方案进行了系统综述[14]；Guo 等人研究了联邦学习环境下的隐私保护标签传播问题[13]；Fu 等人则进一步讨论了局部差分隐私条件下的去中心化社会网络社区检测[12]。此外，Shao 等人提出的 PMCDM 框架直接面向多层网络中的隐私保护多分辨率社区检测问题，是与本课题最接近的代表性研究之一[11]。不过，现有研究仍普遍面临两类难点：一是图结构扰动会直接影响边连接关系，从而降低社区划分精度；二是社区检测过程本身依赖大量邻域统计与目标函数评估，若完全在明文下执行，又会带来隐私泄露风险。现有研究大多在“隐私强度”和“检测效果”之间寻找平衡，但同时兼顾多层网络、可视化展示与系统实现的综合平台仍相对较少，本项目正是在这一背景下展开设计与实现。",
            )
        elif text.startswith("其中 A_ij 表示邻接关系"):
            replace_para_text(
                para,
                "其中 A_ij 表示邻接关系，k_i 和 k_j 为节点度，m 为边数，δ(c_i,c_j) 表示两个节点是否位于同一社区。Louvain 算法采用两阶段贪心思想对模块度进行优化：第一阶段通过逐节点移动寻找局部最优社区归属；第二阶段将当前社区压缩为超级节点构造新图，并继续优化[2]。由于其时间复杂度较低、实现简洁，Louvain 成为本项目的基线算法实现依据。",
            )
        elif text.startswith("单层图只能描述一种关系，而在真实场景中，同一批节点往往存在多种关系。"):
            replace_para_text(
                para,
                "单层图只能描述一种关系，而在真实场景中，同一批节点往往存在多种关系。多层网络通过多个图层共同表示同一节点集合上的不同边关系[4,6]，例如 AUCS 数据集中同时存在 coauthor、facebook、leisure、lunch 和 work 五个关系层。本项目将所有数据集统一整理为“多层图列表”结构，以便用同一实验流程处理真实数据与生成数据。",
            )
        elif text.startswith("差分隐私强调：当任意单条记录发生变化时"):
            replace_para_text(
                para,
                "差分隐私强调：当任意单条记录发生变化时，算法输出结果的分布不应出现显著差异[7-8]。其思想是通过向统计结果中注入受控噪声，使攻击者难以从输出中反推单个个体信息。项目中主要使用拉普拉斯机制对候选边对的得分进行扰动，再根据带噪评分选择保留边，从而实现图拓扑层面的隐私保护。",
            )
        elif text.startswith("同态加密允许在不解密数据的情况下直接进行部分运算。"):
            replace_para_text(
                para,
                "同态加密允许在不解密数据的情况下直接进行部分运算。项目选用 Paillier 加法同态加密体制[9]，对边权及社区内部/外部统计量进行加密模拟。终端侧在上传阶段为边权附加 enc_weight 字段，云服务器可以基于密文执行加法与明文数乘，可信端再对统计值进行解密预览。这样既体现了密文计算思想，也使系统架构更贴近隐私计算场景；同时，这一思路与隐私保护网络结构分析中的已有工作具有一致性[10]。",
            )
        elif text.startswith("为全面验证平台能力，项目接入了多类数据集。"):
            replace_para_text(
                para,
                "为全面验证平台能力，项目接入了多类数据集。Karate Club 用于小规模经典网络验证[1]；AUCS 为多层社交网络数据集，适合测试多层社区检测流程[16-17]；LFR 与 mLFR 用于生成具有真实社区标签的人工网络[3]；BioGRID 则提供真实的生物相互作用多层网络，可用于评估算法在真实复杂网络上的表现[15]。",
            )
        elif text.startswith("在本课题中，隐私预算 ε 与目标函数权重 λ 会同时影响社区检测质量和隐私保护能力。"):
            replace_para_text(
                para,
                "在本课题中，隐私预算 ε 与目标函数权重 λ 会同时影响社区检测质量和隐私保护能力。若采用人工逐组试参的方式，不仅效率较低，而且很难在多个指标之间找到稳定折中。为此，项目在 `src/evolutionary_optimizer.py` 中引入了基于遗传算法的参数优化模块，用于自动搜索更优的参数组合，使平台不仅能够“运行算法”，还能够进一步解释“哪些参数更适合当前数据集和隐私目标”。这一思路也与 PMCDM 等工作中强调的“检测性能与隐私保护权衡优化”目标相一致[11]。",
            )
        elif text.startswith("为验证算法在真实复杂网络中的适用性"):
            replace_para_text(
                para,
                "为验证算法在真实复杂网络中的适用性，项目进一步在 7 个指定 BioGRID 物种网络上执行批量实验。BioGRID 是生物网络分析中常用的真实相互作用数据库，能够提供多种实验系统下的生物互作关系[15]。结果显示，双隐私方法在鸡、白色念珠菌SC5314、非洲爪蟾、人类免疫缺陷病毒1型、大鼠和秀丽隐杆线虫等数据上均能输出合理的社区数，并保持约 0.96 左右的隐私率。",
            )

    # Replace reference block
    start_idx = None
    end_idx = None
    for i, para in enumerate(doc.paragraphs):
        t = para.text.strip()
        if t == "【参考文献】":
            start_idx = i
        elif start_idx is not None and t == "致谢":
            end_idx = i
            break

    if start_idx is None or end_idx is None:
        raise RuntimeError("未找到参考文献区域")

    # remove old reference paragraphs between start and 致谢
    for para in list(doc.paragraphs)[start_idx + 1:end_idx]:
        element = para._element
        element.getparent().remove(element)

    ref_anchor = doc.paragraphs[start_idx]
    for ref in NEW_REFERENCES:
        new_para = ref_anchor.insert_paragraph_before("")
        # move new paragraph after anchor by relocating xml
        element = new_para._element
        parent = element.getparent()
        parent.remove(element)
        ref_anchor._element.addnext(element)
        ref_anchor = new_para
        new_para.add_run(ref)

    doc.save(OUT)


if __name__ == "__main__":
    main()
