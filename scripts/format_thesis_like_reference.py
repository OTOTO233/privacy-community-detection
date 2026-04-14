from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, Cm


ROOT = Path(__file__).resolve().parents[1]
SRC_DOC = ROOT / "output" / "基于隐私保护的社区检测平台设计与实现_论文增强版_补充图表公式_迭代实验.docx"
OUT_DOC = ROOT / "output" / "基于隐私保护的社区检测平台设计与实现_格式对齐版.docx"


TITLE_TEXT = "基于隐私保护的社区检测平台设计与实现"


def set_font(font, west: str | None = None, east: str | None = None, size: float | None = None, bold: bool | None = None):
    if west:
        font.name = west
    if east:
        rPr = font._element.get_or_add_rPr()
        rFonts = rPr.rFonts
        if rFonts is None:
            rFonts = rPr._add_rFonts()
        rFonts.set(qn("w:eastAsia"), east)
    if size is not None:
        font.size = Pt(size)
    if bold is not None:
        font.bold = bold


def ensure_style(doc: Document, name: str):
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    set_font(normal.font, "Times New Roman", "宋体", 10.5, None)
    normal.paragraph_format.first_line_indent = Pt(28)
    normal.paragraph_format.space_before = Pt(2.5)
    normal.paragraph_format.space_after = Pt(2.5)

    s1 = ensure_style(doc, "一级标题")
    set_font(s1.font, None, "黑体", 16, True)
    s1.paragraph_format.space_before = Pt(2.5)
    s1.paragraph_format.space_after = Pt(2.5)

    s2 = ensure_style(doc, "二级标题")
    set_font(s2.font, None, "黑体", 15, True)
    s2.paragraph_format.space_before = Pt(2.5)
    s2.paragraph_format.space_after = Pt(2.5)

    s3 = ensure_style(doc, "三级标题")
    set_font(s3.font, None, "黑体", 14, True)
    s3.paragraph_format.space_before = Pt(2.5)
    s3.paragraph_format.space_after = Pt(2.5)

    abstract = ensure_style(doc, "摘要")
    set_font(abstract.font, "Calibri", "楷体", 12, True)

    caption = ensure_style(doc, "Caption")
    set_font(caption.font, "Arial", "黑体", 10, None)

    ref = ensure_style(doc, "参考文献")
    set_font(ref.font, None, "楷体", 12, True)
    ref.paragraph_format.space_before = Pt(7.8)
    ref.paragraph_format.space_after = Pt(7.8)

    ack = ensure_style(doc, "致谢")
    set_font(ack.font, None, "黑体", 12, True)
    ack.paragraph_format.space_before = Pt(7.8)
    ack.paragraph_format.space_after = Pt(7.8)

    title_style = ensure_style(doc, "毕设标题")
    set_font(title_style.font, "STZhongsong", "华文中宋", 18, True)
    title_style.paragraph_format.space_before = Pt(7.8)
    title_style.paragraph_format.space_after = Pt(7.8)

    body_kaiti = ensure_style(doc, "五号楷体")
    set_font(body_kaiti.font, None, "楷体", 10.5, None)

    abs_en = ensure_style(doc, "abstract内容")
    set_font(abs_en.font, "Times New Roman", "Times New Roman", 12, True)
    abs_en.paragraph_format.space_before = Pt(7.8)
    abs_en.paragraph_format.space_after = Pt(7.8)

    for name, left in [("toc 1", 0), ("toc 2", 11), ("toc 3", 22)]:
        try:
            st = doc.styles[name]
        except KeyError:
            continue
        set_font(st.font, "Times New Roman", "宋体", 10, None)
        st.paragraph_format.left_indent = Pt(left)
        st.paragraph_format.space_before = Pt(0)
        st.paragraph_format.space_after = Pt(0)


def set_runs_font(paragraph, west: str | None = None, east: str | None = None, size: float | None = None, bold: bool | None = None):
    for run in paragraph.runs:
        set_font(run.font, west, east, size, bold)


def is_body_paragraph(p) -> bool:
    txt = p.text.strip()
    if not txt:
        return False
    if p.style.name in {"一级标题", "二级标题", "三级标题", "Caption", "参考文献", "致谢", "toc 1", "toc 2", "toc 3", "毕设标题"}:
        return False
    if txt.startswith("图") or txt.startswith("表"):
        return False
    return True


def format_document(doc: Document) -> None:
    for sec in doc.sections:
        sec.top_margin = Cm(2.54)
        sec.bottom_margin = Cm(2.54)
        sec.left_margin = Cm(3.17)
        sec.right_margin = Cm(3.17)
        for p in sec.header.paragraphs:
            p.text = ""
        for p in sec.footer.paragraphs:
            p.text = ""

    configure_styles(doc)

    started_body = False
    in_ref = False
    in_ack = False
    in_en_abs = False

    for p in doc.paragraphs:
        txt = p.text.strip()
        if not txt:
            continue

        if txt == TITLE_TEXT:
            p.style = doc.styles["毕设标题"]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_runs_font(p, "STZhongsong", "华文中宋", 18, True)
            continue

        if txt.startswith("【参考文献】"):
            p.style = doc.styles["参考文献"]
            in_ref = True
            in_ack = False
            in_en_abs = False
            set_runs_font(p, None, "楷体", 12, True)
            continue

        if txt == "致谢":
            p.style = doc.styles["致谢"]
            in_ref = False
            in_ack = True
            in_en_abs = False
            set_runs_font(p, None, "黑体", 12, True)
            continue

        if txt.startswith("【Abstract】"):
            in_ref = False
            in_ack = False
            in_en_abs = True
            p.style = doc.styles["abstract内容"]
            p.paragraph_format.first_line_indent = Pt(21)
            p.paragraph_format.space_before = Pt(7.8)
            p.paragraph_format.space_after = Pt(7.8)
            set_runs_font(p, "Times New Roman", "Times New Roman", 12, None)
            continue

        if txt.startswith("【摘要】"):
            p.style = doc.styles["五号楷体"]
            p.paragraph_format.first_line_indent = Pt(21)
            p.paragraph_format.space_before = Pt(7.8)
            p.paragraph_format.space_after = Pt(7.8)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_runs_font(p, "Times New Roman", "楷体", 10.5, None)
            continue

        if txt.startswith("【关键词】"):
            p.style = doc.styles["Normal"]
            p.paragraph_format.first_line_indent = Pt(21)
            p.paragraph_format.space_before = Pt(7.8)
            p.paragraph_format.space_after = Pt(7.8)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_runs_font(p, "Times New Roman", "宋体", 10.5, None)
            continue

        if txt.startswith("目 录"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_runs_font(p, "Times New Roman", "宋体", 10.5, None)
            continue

        if p.style.name in {"一级标题", "二级标题", "三级标题"}:
            if p.style.name == "一级标题":
                set_runs_font(p, None, "黑体", 16, True)
            elif p.style.name == "二级标题":
                set_runs_font(p, None, "黑体", 15, True)
            else:
                set_runs_font(p, None, "黑体", 14, True)
            p.paragraph_format.space_before = Pt(2.5)
            p.paragraph_format.space_after = Pt(2.5)
            started_body = True
            in_ref = False
            in_ack = False
            continue

        if p.style.name.startswith("toc"):
            set_runs_font(p, "Times New Roman", "宋体", 10, None)
            continue

        if txt.startswith("学号：") or ("专业" in txt and "学院" in txt):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(7.8)
            p.paragraph_format.space_after = Pt(7.8)
            set_runs_font(p, "Times New Roman", "宋体", 10.5, None)
            continue

        if in_ref:
            p.style = doc.styles["参考文献"]
            p.paragraph_format.left_indent = Pt(21)
            p.paragraph_format.first_line_indent = Pt(-21)
            p.paragraph_format.space_before = Pt(7.8)
            p.paragraph_format.space_after = Pt(7.8)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_runs_font(p, None, "楷体", 10.5, None)
            continue

        if in_ack:
            p.style = doc.styles["Normal"]
            p.paragraph_format.first_line_indent = Pt(21)
            p.paragraph_format.space_before = Pt(7.8)
            p.paragraph_format.space_after = Pt(7.8)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_runs_font(p, "Times New Roman", "宋体", 10.5, None)
            continue

        if in_en_abs:
            p.style = doc.styles["abstract内容"]
            p.paragraph_format.first_line_indent = Pt(21)
            p.paragraph_format.space_before = Pt(7.8)
            p.paragraph_format.space_after = Pt(7.8)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_runs_font(p, "Times New Roman", "Times New Roman", 12, None)
            continue

        if is_body_paragraph(p):
            p.style = doc.styles["Normal"]
            p.paragraph_format.first_line_indent = Pt(21)
            p.paragraph_format.space_before = Pt(7.8)
            p.paragraph_format.space_after = Pt(7.8)
            if p.alignment is None:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_runs_font(p, "Times New Roman", "宋体", 10.5, None)


def main() -> None:
    doc = Document(str(SRC_DOC))
    format_document(doc)
    doc.save(str(OUT_DOC))
    print(OUT_DOC)


if __name__ == "__main__":
    main()
