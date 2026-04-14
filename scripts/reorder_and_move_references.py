from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "output" / "thesis_cn_visual_final_referenced.docx"
OUT = ROOT / "output" / "thesis_cn_visual_final_refs_ordered.docx"

REFERENCE_TEXTS = {
    1: "Fortunato S. Community detection in graphs[J]. Physics Reports, 2010, 486(3-5): 75-174.",
    2: "Blondel V D, Guillaume J L, Lambiotte R, et al. Fast unfolding of communities in large networks[J]. Journal of Statistical Mechanics: Theory and Experiment, 2008(10): P10008.",
    3: "Lancichinetti A, Fortunato S, Radicchi F. Benchmark graphs for testing community detection algorithms[J]. Physical Review E, 2008, 78(4): 046110.",
    4: "Mucha P J, Richardson T, Macon K, et al. Community structure in time-dependent, multiscale, and multiplex networks[J]. Science, 2010, 328(5980): 876-878.",
    5: "Li J, Lai S, Shuai Z, et al. A comprehensive review of community detection in graphs[J]. Neurocomputing, 2024, 600: 128169.",
    6: "Kivelä M, Arenas A, Barthélemy M, et al. Multilayer networks[J]. Journal of Complex Networks, 2014, 2(3): 203-271.",
    7: "Zhu T, Li G, Zhou W, et al. Differentially private data publishing and analysis: A survey[J]. IEEE Transactions on Knowledge and Data Engineering, 2017, 29(8): 1619-1638.",
    8: "Dwork C, Roth A. The Algorithmic Foundations of Differential Privacy[J]. Foundations and Trends in Theoretical Computer Science, 2014, 9(3-4): 211-407.",
    9: "Paillier P. Public-key cryptosystems based on composite degree residuosity classes[C]// Advances in Cryptology—EUROCRYPT'99. Berlin: Springer, 1999: 223-238.",
    10: "Ma L, Huang X, Li J, et al. Privacy-preserving global structural balance computation in signed networks[J]. IEEE Transactions on Computational Social Systems, 2019, 7(1): 164-177.",
    11: "Shao Z, Ma L, Lin Q, et al. PMCDM: Privacy-preserving multiresolution community detection in multiplex networks[J]. Knowledge-Based Systems, 2022, 244: 108542.",
    12: "Fu N, Ni W, Hou L, et al. Community detection in decentralized social networks with local differential privacy[J]. Information Sciences, 2024, 661: 120164.",
    13: "Guo K, Chen D, Huang Q, et al. Privacy-preserving multi-label propagation based on federated learning[J]. IEEE Transactions on Network Science and Engineering, 2023, 11(1): 886-899.",
    14: "Ferrag M A, Maglaras L, Ahmim A. Privacy-preserving schemes for ad hoc social networks: A survey[J]. IEEE Communications Surveys & Tutorials, 2017, 19(4): 3015-3045.",
    15: "Oughtred R, Rust J, Chang C, et al. The BioGRID interaction database: 2019 update[J]. Nucleic Acids Research, 2019, 47(D1): D529-D541.",
    16: "Magnani M, Micenková B, Rossi L. Combinatorial analysis of multiple networks[EB/OL]. arXiv, 2013: arXiv:1303.4986.",
    17: "Bródka P. A Method for Group Extraction and Analysis in Multilayer Social Networks[EB/OL]. arXiv, 2016: arXiv:1612.02377.",
}

CITE_RE = re.compile(r"\[(\d+(?:-\d+)?(?:,\d+(?:-\d+)?)*)\]")
REF_LINE_RE = re.compile(r"^\[(\d+)\]")


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def insert_paragraph_before(paragraph: Paragraph, text: str, style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    new_para.add_run(text)
    return new_para


def insert_paragraph_after(paragraph: Paragraph, text: str, style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    new_para.add_run(text)
    return new_para


def expand_token(token: str) -> list[int]:
    token = token.strip()
    if "-" in token:
        start, end = token.split("-", 1)
        return list(range(int(start), int(end) + 1))
    return [int(token)]


def compress_numbers(nums: list[int]) -> str:
    nums = sorted(dict.fromkeys(nums))
    if not nums:
        return ""
    ranges: list[str] = []
    start = prev = nums[0]
    for n in nums[1:]:
        if n == prev + 1:
            prev = n
        else:
            ranges.append(f"{start}-{prev}" if start != prev else str(start))
            start = prev = n
    ranges.append(f"{start}-{prev}" if start != prev else str(start))
    return ",".join(ranges)


def reorder_citations(text: str, mapping: dict[int, int]) -> str:
    def repl(match: re.Match[str]) -> str:
        items = []
        for token in match.group(1).split(","):
            items.extend(expand_token(token))
        mapped = [mapping.get(n, n) for n in items]
        return f"[{compress_numbers(mapped)}]"

    return CITE_RE.sub(repl, text)


def main() -> None:
    doc = Document(SRC)
    paragraphs = list(doc.paragraphs)

    # Identify last final acknowledgement heading
    thanks_indices = [i for i, p in enumerate(paragraphs) if p.text.strip() == "致谢"]
    if not thanks_indices:
        raise RuntimeError("未找到致谢标题")
    last_thanks_idx = thanks_indices[-1]

    # Remove all existing reference headings and reference lines
    for para in list(doc.paragraphs):
        text = para.text.strip()
        if text == "【参考文献】" or REF_LINE_RE.match(text):
            remove_paragraph(para)

    # Remove earlier orphan 致谢 headings before final thanks
    paragraphs = list(doc.paragraphs)
    thanks_indices = [i for i, p in enumerate(paragraphs) if p.text.strip() == "致谢"]
    if not thanks_indices:
        raise RuntimeError("清理后未找到致谢标题")
    last_thanks_idx = thanks_indices[-1]
    for i in reversed(thanks_indices[:-1]):
        remove_paragraph(doc.paragraphs[i])

    paragraphs = list(doc.paragraphs)
    thanks_para = next(p for p in paragraphs if p.text.strip() == "致谢")

    # Body paragraphs used for citation ordering: before final thanks, excluding abstract english tail if any
    current_paragraphs = list(doc.paragraphs)
    thanks_index = next(i for i, p in enumerate(current_paragraphs) if p.text.strip() == "致谢")
    cited_order: list[int] = []
    seen: set[int] = set()
    for para in current_paragraphs[:thanks_index]:
        text = para.text.strip()
        if not text:
            continue
        if text in {"【参考文献】", "致谢"} or text.startswith("["):
            continue
        for match in CITE_RE.finditer(text):
            nums: list[int] = []
            for token in match.group(1).split(","):
                nums.extend(expand_token(token))
            for num in nums:
                if num in REFERENCE_TEXTS and num not in seen:
                    seen.add(num)
                    cited_order.append(num)

    final_order = cited_order + [n for n in sorted(REFERENCE_TEXTS) if n not in seen]
    mapping = {old: idx + 1 for idx, old in enumerate(final_order)}

    # Update citations in all non-reference paragraphs
    for para in doc.paragraphs:
        text = para.text
        stripped = text.strip()
        if not stripped or stripped == "【参考文献】" or stripped == "致谢" or REF_LINE_RE.match(stripped):
            continue
        new_text = reorder_citations(text, mapping)
        if new_text != text:
            para.clear()
            para.add_run(new_text)

    # Insert final reference section before last acknowledgement
    heading_para = insert_paragraph_before(thanks_para, "【参考文献】", style="参考文献")
    anchor = heading_para
    ordered_refs = [REFERENCE_TEXTS[old] for old in final_order]
    for idx, ref_text in enumerate(ordered_refs, start=1):
        anchor = insert_paragraph_after(anchor, f"[{idx}] {ref_text}")

    doc.save(OUT)


if __name__ == "__main__":
    main()
