from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "output" / "thesis_latex_layout_fixed.docx"
OUT = ROOT / "output" / "thesis_cn_visual_final.docx"
OUTPUT = ROOT / "output"


PARAGRAPH_REPLACE = {
    "表5-2 Karate Club 数据集算法对比结果": "表5-2 空手道俱乐部数据集算法对比结果",
    "表5-3 AUCS 数据集算法对比结果": "表5-3 多层社交网络数据集算法对比结果",
    "表5-4 指定 7 个 BioGRID 物种的 DH-Louvain 结果": "表5-4 七个真实生物互作网络物种的双隐私方法结果",
    "图5-4 DH-Louvain 在 7 个 BioGRID 物种上的指标折线图": "图5-4 双隐私方法在七个真实生物互作网络物种上的指标折线图",
    "图5-4-2 五个真实 BioGRID 网络在不同迭代次数下的 D/n、Q、NMI 折线图": "图5-4-2 五个真实生物网络在不同迭代次数下的归一化模块密度、模块度和归一化互信息折线图",
    "图5-5 LFR 参数控制实验中 Q 随 μ 变化的折线图": "图5-5 基准网络参数控制实验中模块度随 μ 变化的折线图",
    "图5-6 LFR 参数控制实验中 D 随 μ 变化的折线图": "图5-6 基准网络参数控制实验中模块密度随 μ 变化的折线图",
    "图5-7 LFR 参数控制实验中 NMI 随 μ 变化的折线图": "图5-7 基准网络参数控制实验中归一化互信息随 μ 变化的折线图",
    "表5-5 Karate Club 数据集参数优化前后对比": "表5-5 空手道俱乐部数据集参数优化前后对比",
    "图5-8 遗传算法参数优化过程折线图": "图5-8 遗传算法参数优化过程折线图",
    "图5-9 进化算法优化前后平均指标对比图": "图5-9 进化算法优化前后平均指标对比图",
    "图5-10 Karate Club 数据集的 DH-Louvain 可视化结果": "图5-10 空手道俱乐部数据集的双隐私方法可视化结果",
    "图5-11 AUCS 数据集的社区检测可视化结果": "图5-11 多层社交网络数据集的社区检测可视化结果",
}

TEXT_REPLACE = {
    "Karate Club": "空手道俱乐部",
    "DH-Louvain": "双隐私方法",
    "S-Louvain": "标准方法",
    "PD-Louvain": "加密方法",
    "R-Louvain": "随机扰动方法",
    "DP-Louvain": "差分隐私方法",
    "K-Louvain": "度匿名方法",
    "Gallus gallus": "鸡",
    "Bos taurus": "牛",
    "Candida albicans SC5314": "白色念珠菌SC5314",
    "Xenopus laevis": "非洲爪蟾",
    "Human Immunodeficiency Virus 1": "人类免疫缺陷病毒1型",
    "Rattus norvegicus": "大鼠",
    "Caenorhabditis elegans": "秀丽隐杆线虫",
    "Plasmodium": "恶性疟原虫3D7",
    "HIV-1": "人类免疫缺陷病毒1型",
    "E.coli": "大肠杆菌K12 MG1655",
    "Xenopus": "非洲爪蟾",
}

CELL_REPLACE = {
    "算法": "算法",
    "Q": "模块度Q",
    "D": "模块密度D",
    "NMI": "归一化互信息",
    "pr": "隐私率",
    "fitness": "适应度",
    "Karate、AUCS、LFR、mLFR、BioGRID": "空手道俱乐部、多层社交网络、LFR基准网络、mLFR基准网络、生物互作数据库",
    "S-Louvain": "标准方法",
    "PD-Louvain": "加密方法",
    "R-Louvain": "随机扰动方法",
    "DP-Louvain": "差分隐私方法",
    "K-Louvain": "度匿名方法",
    "DH-Louvain": "双隐私方法",
    "Gallus gallus": "鸡",
    "Bos taurus": "牛",
    "Candida albicans SC5314": "白色念珠菌SC5314",
    "Xenopus laevis": "非洲爪蟾",
    "Human Immunodeficiency Virus 1": "人类免疫缺陷病毒1型",
    "Rattus norvegicus": "大鼠",
    "Caenorhabditis elegans": "秀丽隐杆线虫",
}

FIGURE_MAP = {
    "图5-4 双隐私方法在七个真实生物互作网络物种上的指标折线图": OUTPUT / "chart_biogrid_dh_species_lines_cn.png",
    "图5-4-2 五个真实生物网络在不同迭代次数下的归一化模块密度、模块度和归一化互信息折线图": OUTPUT / "chart_biogrid_5_iterations_triptych_cn.png",
    "图5-5 基准网络参数控制实验中模块度随 μ 变化的折线图": OUTPUT / "chart_lfr_mu_group_modularity_n300_all_cn.png",
    "图5-6 基准网络参数控制实验中模块密度随 μ 变化的折线图": OUTPUT / "chart_lfr_mu_group_density_n300_all_cn.png",
    "图5-7 基准网络参数控制实验中归一化互信息随 μ 变化的折线图": OUTPUT / "chart_lfr_mu_group_nmi_n300_all_cn.png",
    "图5-8 遗传算法参数优化过程折线图": OUTPUT / "chart_ea_fitness_history_cn.png",
    "图5-9 进化算法优化前后平均指标对比图": OUTPUT / "chart_ea_before_after_compare_cn.png",
    "图5-10 空手道俱乐部数据集的双隐私方法可视化结果": OUTPUT / "karate_dh_vis_cn.png",
    "图5-11 多层社交网络数据集的社区检测可视化结果": OUTPUT / "aucs_dh_vis_cn.png",
    "图5-12 各算法在检测能力与隐私保护上的综合排名图": OUTPUT / "chart_conclusion_algorithm_ranking_cn.png",
}

NEW_BIOGRID_FIGURES = [
    ("图5-11-1 恶性疟原虫3D7数据集的社区可视化结果", OUTPUT / "biogrid_plasmodium_vis_cn.png"),
    ("图5-11-2 恶性疟原虫3D7数据集的多层网络可视化结果", OUTPUT / "biogrid_plasmodium_multilayer_cn.png"),
    ("图5-11-3 人类免疫缺陷病毒1型数据集的社区可视化结果", OUTPUT / "biogrid_hiv1_vis_cn.png"),
    ("图5-11-4 人类免疫缺陷病毒1型数据集的多层网络可视化结果", OUTPUT / "biogrid_hiv1_multilayer_cn.png"),
    ("图5-11-5 大肠杆菌K12 MG1655数据集的社区可视化结果", OUTPUT / "biogrid_ecoli_vis_cn.png"),
    ("图5-11-6 大肠杆菌K12 MG1655数据集的多层网络可视化结果", OUTPUT / "biogrid_ecoli_multilayer_cn.png"),
]


def paragraph_contains_drawing(paragraph: Paragraph) -> bool:
    return bool(paragraph._element.xpath(".//w:drawing"))


def insert_paragraph_after(paragraph: Paragraph, style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    return new_para


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    parent.remove(element)


def replace_text(text: str) -> str:
    updated = text
    for old, new in TEXT_REPLACE.items():
        updated = updated.replace(old, new)
    return updated


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def add_picture_before_caption(caption_para: Paragraph, image_path: Path, width_cm: float = 15.0) -> None:
    previous = caption_para._p.getprevious()
    if previous is not None:
        prev_para = Paragraph(previous, caption_para._parent)
        if paragraph_contains_drawing(prev_para):
            remove_paragraph(prev_para)
    pic_para = OxmlElement("w:p")
    caption_para._p.addprevious(pic_para)
    picture_paragraph = Paragraph(pic_para, caption_para._parent)
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.add_run().add_picture(str(image_path), width=Cm(width_cm))


def add_caption_after(paragraph: Paragraph, caption: str) -> Paragraph:
    caption_para = insert_paragraph_after(paragraph)
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_para.style = "Caption"
    caption_para.add_run(caption)
    return caption_para


def add_picture_after(paragraph: Paragraph, image_path: Path, caption: str, width_cm: float = 15.0) -> Paragraph:
    pic_para = insert_paragraph_after(paragraph)
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_para.add_run().add_picture(str(image_path), width=Cm(width_cm))
    return add_caption_after(pic_para, caption)


def main() -> None:
    doc = Document(SRC)

    for paragraph in doc.paragraphs:
        raw = paragraph.text.strip()
        if raw in PARAGRAPH_REPLACE:
            set_paragraph_text(paragraph, PARAGRAPH_REPLACE[raw])
            raw = paragraph.text.strip()
        else:
            updated = replace_text(paragraph.text)
            if updated != paragraph.text:
                set_paragraph_text(paragraph, updated)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                new_text = CELL_REPLACE.get(text, replace_text(text))
                if new_text != text:
                    cell.text = new_text

    for paragraph in doc.paragraphs:
        caption = paragraph.text.strip()
        if caption in FIGURE_MAP:
            add_picture_before_caption(paragraph, FIGURE_MAP[caption], width_cm=15.0 if "图5-11" not in caption and "图5-10" not in caption else 14.5)

    anchor = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "图5-11 多层社交网络数据集的社区检测可视化结果":
            anchor = paragraph
            break

    if anchor is not None:
        intro = insert_paragraph_after(anchor)
        intro.style = "Normal"
        intro.add_run(
            "为了进一步展示真实生物网络中的社区结构，本文补充给出了前三个真实数据集的社区可视化结果和多层网络可视化结果。"
            "从这些图可以看到，双隐私方法在恶性疟原虫3D7、人类免疫缺陷病毒1型和大肠杆菌K12 MG1655数据上都能够形成较明显的聚类结构；"
            "同时，多层可视化结果也说明不同实验系统层之间存在既共享又差异化的连接模式，这与多层网络建模的设定是一致的。"
        )
        current = intro
        for caption, image_path in NEW_BIOGRID_FIGURES:
            current = add_picture_after(current, image_path, caption, width_cm=15.2)
        outro = insert_paragraph_after(current)
        outro.style = "Normal"
        outro.add_run(
            "综合这六张补充图可以进一步说明：对于真实生物互作网络，双隐私方法不仅能够在聚合图上给出可解释的社区划分，"
            "还能够在多层结构视角下保留不同层之间的拓扑差异，因此平台的可视化结果与前文的定量实验结论能够相互印证。"
        )

    doc.save(OUT)


if __name__ == "__main__":
    main()
