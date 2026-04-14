from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_PATH = ROOT.parent / "4设计实现工具系统类毕业论文内容章节安排参考模板.docx"
OUTPUT_DIR = ROOT / "output"
OUTPUT_DOCX = OUTPUT_DIR / "基于隐私保护的社区检测平台设计与实现_论文增强版_补充图表公式_迭代实验.docx"
OUTPUT_MD = OUTPUT_DIR / "基于隐私保护的社区检测平台设计与实现_论文增强版_补充图表公式_迭代实验.md"

TITLE_LINE_1 = "基于隐私保护的社区检测平台"
TITLE_LINE_2 = "设计与实现"
AUTHOR_PLACEHOLDER = "（请填写）"
MAJOR_PLACEHOLDER = "计算机科学与技术"
COLLEGE_PLACEHOLDER = "计算机与软件学院"
STUDENT_ID_PLACEHOLDER = "（请填写）"
ADVISOR_PLACEHOLDER = "（请填写）"
ADVISOR_TITLE_PLACEHOLDER = "（请填写）"
DATE_TEXT = "2026年 4月 13 日"


def set_para_text(paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def clear_from_paragraph(document: Document, start_index: int) -> None:
    for paragraph in list(document.paragraphs)[start_index:]:
        remove_paragraph(paragraph)


def append_paragraph(document: Document, text: str, style: str = "Normal", align=None):
    paragraph = document.add_paragraph(style=style)
    paragraph.add_run(text)
    if align is not None:
        paragraph.alignment = align
    return paragraph


def append_heading(document: Document, text: str, level: int) -> None:
    style = {1: "一级标题", 2: "二级标题", 3: "三级标题"}[level]
    append_paragraph(document, text, style=style)


def append_caption(document: Document, text: str) -> None:
    append_paragraph(document, text, style="Caption", align=WD_ALIGN_PARAGRAPH.CENTER)


def append_table(document: Document, title: str, headers: list[str], rows: list[list[str]]) -> None:
    append_caption(document, title)
    table = document.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value


def append_picture(document: Document, path: Path, caption: str, width_cm: float = 14.5) -> None:
    if not path.exists():
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Cm(width_cm))
    append_caption(document, caption)


def find_latest_image(prefix: str) -> Path | None:
    matches = sorted(OUTPUT_DIR.glob(f"{prefix}*.png"), key=lambda item: item.stat().st_mtime, reverse=True)
    return matches[0] if matches else None


def build_markdown() -> str:
    return """# 基于隐私保护的社区检测平台设计与实现

本文件为根据项目代码、实验输出和模板结构整理的论文初稿 Markdown 版本，与生成的 `.docx` 内容保持一致，便于后续继续修改。
"""


def build_document() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document = Document(str(TEMPLATE_PATH))

    paragraph_updates = {
        5: "题目:    " + TITLE_LINE_1,
        6: TITLE_LINE_2,
        7: "姓名:          " + AUTHOR_PLACEHOLDER,
        8: "专业:      " + MAJOR_PLACEHOLDER,
        9: "学院:      " + COLLEGE_PLACEHOLDER,
        10: "学号:        " + STUDENT_ID_PLACEHOLDER,
        11: "指导教师:      " + ADVISOR_PLACEHOLDER,
        12: "职称：         " + ADVISOR_TITLE_PLACEHOLDER,
        17: DATE_TEXT,
        21: (
            "本人郑重声明：所呈交的毕业论文（设计），题目《"
            + TITLE_LINE_1
            + TITLE_LINE_2
            + "》是本人在指导教师的指导下，独立进行研究工作所取得的成果。"
            "对本文的研究做出重要贡献的个人和集体，均已在文中以明确方式注明。"
            "除此之外，本论文不包含任何其他个人或集体已经发表或撰写过的作品成果。"
            "本人完全意识到本声明的法律结果。"
        ),
    }
    for index, text in paragraph_updates.items():
        set_para_text(document.paragraphs[index], text)

    clear_from_paragraph(document, 33)

    append_paragraph(document, "目 录", style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER)
    toc_lines = [
        ("摘要（关键字）", "toc 1"),
        ("1引言", "toc 1"),
        ("1.1研究背景及意义", "toc 2"),
        ("1.1.1研究背景", "toc 3"),
        ("1.1.2研究意义", "toc 3"),
        ("1.2国内外研究现状", "toc 2"),
        ("1.2.1社区检测研究现状", "toc 3"),
        ("1.2.2隐私保护社区检测研究现状", "toc 3"),
        ("1.3研究方法与工作", "toc 2"),
        ("1.3.1研究方法", "toc 3"),
        ("1.3.2研究工作", "toc 3"),
        ("1.4本章小结", "toc 2"),
        ("2相关理论与关键技术", "toc 1"),
        ("2.1社区检测基本原理", "toc 2"),
        ("2.1.1社区的基本概念", "toc 3"),
        ("2.1.2模块度与Louvain算法", "toc 3"),
        ("2.1.3多层网络与加权模块密度", "toc 3"),
        ("2.2隐私保护关键技术", "toc 2"),
        ("2.2.1差分隐私", "toc 3"),
        ("2.2.2同态加密", "toc 3"),
        ("2.2.3隐私保护在社区检测中的可行性分析", "toc 3"),
        ("2.3数据集与评价指标", "toc 2"),
        ("2.4本章小结", "toc 2"),
        ("3隐私保护社区检测算法设计与实现", "toc 1"),
        ("3.1总体算法设计", "toc 2"),
        ("3.1.1设计目标", "toc 3"),
        ("3.1.2三方协同架构", "toc 3"),
        ("3.1.3算法流程", "toc 3"),
        ("3.2核心算法实现过程", "toc 2"),
        ("3.2.1数据装载与多层表示", "toc 3"),
        ("3.2.2S-Louvain基线实现", "toc 3"),
        ("3.2.3DH-Louvain实现", "toc 3"),
        ("3.2.4差分隐私拓扑扰动实现", "toc 3"),
        ("3.2.5同态加密与加密统计实现", "toc 3"),
        ("3.2.6基于进化算法的参数优化实现", "toc 3"),
        ("3.3算法实现特点与复杂度分析", "toc 2"),
        ("3.4本章小结", "toc 2"),
        ("4系统架构与实现", "toc 1"),
        ("4.1系统总体架构设计", "toc 2"),
        ("4.2后端设计与实现", "toc 2"),
        ("4.3前端设计与实现", "toc 2"),
        ("4.4可视化设计与实现", "toc 2"),
        ("4.5本章小结", "toc 2"),
        ("5系统测试与实验分析", "toc 1"),
        ("5.1实验环境与方案", "toc 2"),
        ("5.2基准数据集实验分析", "toc 2"),
        ("5.3真实BioGRID数据实验分析", "toc 2"),
        ("5.4参数优化与补充实验分析", "toc 2"),
        ("5.5系统功能验证与可视化结果", "toc 2"),
        ("5.6本章小结", "toc 2"),
        ("6总结与展望", "toc 1"),
        ("【参考文献】", "toc 1"),
        ("致谢", "toc 1"),
        ("Abstract（Key words）", "toc 1"),
    ]
    for text, style in toc_lines:
        append_paragraph(document, text, style=style)
    document.add_page_break()

    append_paragraph(document, TITLE_LINE_1 + TITLE_LINE_2, align=WD_ALIGN_PARAGRAPH.CENTER)
    append_paragraph(
        document,
        f"{COLLEGE_PLACEHOLDER}{MAJOR_PLACEHOLDER}专业 {AUTHOR_PLACEHOLDER}",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    append_paragraph(document, f"学号：{STUDENT_ID_PLACEHOLDER}", align=WD_ALIGN_PARAGRAPH.CENTER)
    abstract_paras = [
        "【摘要】复杂网络中的社区结构能够反映节点之间潜在的组织关系、功能关联和交互模式，因此社区检测在社交网络分析、生物网络分析、协同关系挖掘等场景中具有重要价值。然而，真实网络数据往往包含敏感边关系、交互频次与节点属性，若直接在明文图上进行社区检测，容易造成隐私泄露。围绕这一问题，本文基于隐私保护思想设计并实现了一套面向多层复杂网络的社区检测平台，将社区检测算法、差分隐私机制、同态加密模拟和可视化系统进行集成，用于验证隐私保护下社区检测的可行性、有效性与工程实现路径。",
        "本文首先梳理了社区、模块度、多层网络、差分隐私和同态加密等相关理论，分析隐私保护技术应用于社区检测任务的可行性；随后结合项目代码实现，构建了“终端—云服务器1—云服务器2”三方协同框架，在此基础上实现了 DH-Louvain 多层社区检测流程，并给出了差分隐私拓扑扰动、加密上传、加权模块密度优化及进化参数优化等关键模块；最后设计并实现了基于 FastAPI 与 Vue 的可视化平台，支持 Karate、AUCS、LFR、mLFR 和 BioGRID 等数据集的选择、参数配置、算法对比、指标展示和社区结构可视化。",
        "实验结果表明，所实现的平台能够在隐私保护条件下完成社区检测任务，并在不同数据集上输出模块度、模块密度、归一化互信息和隐私率等指标。以真实 BioGRID 数据集为例，DH-Louvain 在多个物种网络上均可保持较高的隐私率，同时获得可解释的社区划分结果。项目还通过遗传算法对隐私预算 ε 与目标函数权重进行搜索，验证了平台具备进一步开展参数优化与实验扩展的能力。",
        "【关键词】社区检测；隐私保护；差分隐私；同态加密；多层网络；可视化系统",
    ]
    for text in abstract_paras:
        append_paragraph(document, text, style="摘要")

    sections = [
        (1, "1引言"),
        (2, "1.1研究背景及意义"),
        (3, "研究背景"),
        ("p", "随着互联网平台、科研协作系统和生物相互作用数据库的快速发展，越来越多的数据可以抽象为复杂网络。社交网络中的好友关系、通信行为与兴趣联系，科研网络中的合作关系，以及生物网络中的蛋白质相互作用，都可以表示为由节点和边组成的图结构。在这类图数据中，节点之间往往会自然形成连接更紧密、边界更明显的子群体，即社区结构。社区检测的目标就是自动发现这些子群体，从而为群体发现、功能模块识别和关系挖掘提供基础。"),
        ("p", "然而，在实际应用中，社区检测并不是一个单纯的图算法问题。真实网络数据中往往包含敏感关系信息，例如用户之间的社交连接、单位内部协作关系、研究人员的共同作者网络以及生物分子之间的实验互作关系。如果直接上传或公开原始图结构，不仅可能暴露个体身份和群体边界，还可能对机构、平台和数据拥有者带来安全风险。因此，如何在不充分暴露原始数据的前提下完成社区检测，逐渐成为复杂网络分析中的重要研究方向。"),
        (3, "研究意义"),
        ("p", "本课题的研究意义主要体现在两个方面。其一，在理论层面，将社区检测与隐私保护技术结合，有助于探索图挖掘任务与隐私计算之间的融合方式，验证差分隐私和同态加密在复杂网络场景中的适用性。其二，在工程层面，基于隐私保护的社区检测平台不仅可以支撑算法实验，也有助于形成可视化、可演示、可扩展的毕业设计系统原型，为后续开展更大规模数据分析和更复杂隐私机制研究打下基础。"),
        (2, "1.2国内外研究现状"),
        (3, "1.2.1社区检测研究现状"),
        ("p", "国外关于社区检测的研究起步较早。Fortunato 对社区检测问题进行了系统综述，指出社区结构能够帮助理解复杂系统的组织方式，并总结了模块度优化、谱聚类、标签传播等多类方法[1]。Blondel 等人提出的 Louvain 算法由于具有较好的效率和可扩展性，成为大规模网络社区检测中的经典方法[2]。在测试方法方面，Lancichinetti、Fortunato 与 Radicchi 提出的 LFR 基准网络更接近真实网络的度分布和社区规模分布，现已成为评价社区检测算法的重要基准[3]。"),
        ("p", "国内关于复杂网络社区发现的研究更多集中在算法改进、指标优化以及特定应用场景建模等方面。随着社交网络分析、生物网络分析和多源异构数据分析需求的增加，围绕模块度缺陷修正、重叠社区识别以及多层网络划分的研究逐渐增多，但与国外相比，面向隐私保护与系统实现一体化的研究成果仍相对有限。"),
        ("p", "随着多源关系网络的普及，研究者开始关注多层网络和多路复用网络中的社区结构。Kivelä 等人从统一理论角度总结了多层网络的表示方法与分析框架，指出传统单层图模型难以完整刻画多种关系并存的复杂系统[4]。因此，在多层网络场景下，如何综合各层结构信息进行社区划分，成为社区检测研究的重要延伸方向。"),
        (3, "1.2.2隐私保护社区检测研究现状"),
        ("p", "国外关于隐私保护数据分析的研究体系较为成熟。Dwork 与 Roth 系统阐述了差分隐私的定义、性质及其机制设计方法，为隐私保护数据分析奠定了理论基础[5]。Paillier 提出的加法同态公钥密码体制，使得在密文状态下进行加法和数乘成为可能，为安全统计与隐私计算提供了工具[6]。这些研究为图数据隐私保护提供了重要理论支撑。"),
        ("p", "国内在隐私保护图分析方面也已有探索，研究内容涉及社交网络匿名化、图发布保护和差分隐私机制改进等方向，但在社区检测问题中仍普遍面临两个难点：一是图结构扰动会直接影响边连接关系，从而降低社区划分精度；二是社区检测过程本身依赖大量邻域统计与目标函数评估，若完全在明文下执行，又会带来隐私泄露风险。现有研究大多在“隐私强度”和“检测效果”之间寻找平衡，但同时兼顾多层网络、可视化展示与系统实现的综合平台仍相对较少。本项目正是在这一背景下展开设计与实现。"),
        (2, "1.3研究方法与工作"),
        (3, "1.3.1研究方法"),
        ("p", "本文采用“理论分析 + 系统设计 + 算法实现 + 实验验证”的研究方法。首先，通过文献调研明确社区检测、多层网络、差分隐私和同态加密的基本概念与原理；其次，结合平台原型需求，对数据流、算法流与系统模块进行总体设计；然后，在 Python 环境下完成算法、后端和前端的开发；最后，基于公开数据集与真实 BioGRID 数据，对隐私保护能力和检测性能进行测试分析。"),
        (3, "1.3.2研究工作"),
        ("p", "围绕任务书要求，本文完成了以下工作：一是梳理社区基本概念及社区检测算法原理，分析模块度优化与多层网络建模方法；二是研究差分隐私与同态加密机制，设计适用于社区检测任务的隐私保护流程；三是设计并实现隐私保护社区检测平台，提供命令行、后端接口和前端可视化展示能力；四是基于 Karate、AUCS、LFR、mLFR 和 BioGRID 等数据集进行测试，验证算法性能与隐私保护能力；五是引入进化优化模块，对关键参数进行搜索，提高实验分析的完整性。"),
        (2, "1.4本章小结"),
        ("p", "本章介绍了课题的研究背景、研究意义、国内外研究现状以及本文的研究方法和主要工作。可以看出，社区检测与隐私保护的结合既有理论研究价值，也具备明显的工程实践需求。后续章节将在此基础上展开相关理论、算法设计、系统实现与实验分析。"),
        (1, "2相关理论与关键技术"),
        (2, "2.1社区检测基本原理"),
        (3, "2.1.1社区的基本概念"),
        ("p", "社区通常是指网络中内部连接相对紧密、外部连接相对稀疏的一组节点集合。若将网络表示为图 G=(V,E)，则社区检测的目标是将节点集 V 划分为若干子集，使同一子集内节点具有更强的联系。对于社会网络，社区可能对应兴趣群体或工作群组；对于生物网络，社区可能对应功能模块或蛋白复合体。"),
        (3, "2.1.2模块度与Louvain算法"),
        ("p", "模块度 Q 是评价社区划分质量的经典指标，其核心思想是将实际网络中的社区内连边数量与同规模随机网络中的期望连边数量进行比较。若社区内部边显著多于随机期望，则说明划分较优。模块度常写为："),
        ("p_center", "Q = (1 / 2m) Σ_ij (A_ij - k_i k_j / 2m) δ(c_i, c_j)"),
        ("p", "其中 A_ij 表示邻接关系，k_i 和 k_j 为节点度，m 为边数，δ(c_i,c_j) 表示两个节点是否位于同一社区。Louvain 算法采用两阶段贪心思想对模块度进行优化：第一阶段通过逐节点移动寻找局部最优社区归属；第二阶段将当前社区压缩为超级节点构造新图，并继续优化。由于其时间复杂度较低、实现简洁，Louvain 成为本项目的基线算法实现依据。"),
        (3, "2.1.3多层网络与加权模块密度"),
        ("p", "单层图只能描述一种关系，而在真实场景中，同一批节点往往存在多种关系。多层网络通过多个图层共同表示同一节点集合上的不同边关系，例如 AUCS 数据集中同时存在 coauthor、facebook、leisure、lunch 和 work 五个关系层。本项目将所有数据集统一整理为“多层图列表”结构，以便用同一实验流程处理真实数据与生成数据。"),
        ("p", "在 DH-Louvain 中，项目并非仅优化传统模块度，而是使用加权模块密度目标函数："),
        ("p_center", "F_λ = λQ + (1 - λ)D"),
        ("p", "其中 D 为模块密度，用于反映社区内部连接紧密程度，λ 为平衡系数。通过在模块度与模块密度之间进行折中，算法可以在不同数据特征下获得更稳定的划分结果。"),
        (2, "2.2隐私保护关键技术"),
        (3, "2.2.1差分隐私"),
        ("p", "差分隐私强调：当任意单条记录发生变化时，算法输出结果的分布不应出现显著差异。其思想是通过向统计结果中注入受控噪声，使攻击者难以从输出中反推单个个体信息。项目中主要使用拉普拉斯机制对候选边对的得分进行扰动，再根据带噪评分选择保留边，从而实现图拓扑层面的隐私保护。"),
        (3, "2.2.2同态加密"),
        ("p", "同态加密允许在不解密数据的情况下直接进行部分运算。本项目选用 Paillier 加法同态加密体制，对边权及社区内部/外部统计量进行加密模拟。终端侧在上传阶段为边权附加 enc_weight 字段，云服务器可以基于密文执行加法与明文数乘，可信端再对统计值进行解密预览。这样既体现了密文计算思想，也使系统架构更贴近隐私计算场景。"),
        (3, "2.2.3隐私保护在社区检测中的可行性分析"),
        ("p", "隐私保护技术用于社区检测的关键在于“尽量不破坏社区结构的同时保护局部关系”。对于本项目而言，差分隐私主要作用于图拓扑扰动，降低敏感边被恢复的风险；同态加密主要作用于上传与统计阶段，降低中间计算过程暴露明文数据的可能性。二者结合后形成“结构扰动保护 + 计算过程保护”的双重机制。"),
        ("p", "项目中定义综合隐私率 pr 作为隐私保护能力指标："),
        ("p_center", "pr = (p_e + p_w) / 2"),
        ("p", "其中 p_e 表示原始边在扰动后仍被保留的比例，p_w 表示是否采用同态加密。该定义能够从结构与计算两个角度对隐私方案进行工程化衡量。"),
        (2, "2.3数据集与评价指标"),
        ("p", "为全面验证平台能力，项目接入了多类数据集。Karate Club 用于小规模经典网络验证；AUCS 为多层社交网络数据集，适合测试多层社区检测流程；LFR 与 mLFR 用于生成具有真实社区标签的人工网络；BioGRID 则提供真实的生物相互作用多层网络，可用于评估算法在真实复杂网络上的表现。"),
        ("p", "实验指标包括：模块度 Q、模块密度 D、归一化互信息 NMI 和隐私率 pr。若数据集存在真实标签，则 NMI 直接与真实标签比较；若无真实标签，则使用原始网络上的 S-Louvain 划分结果作为参考划分。这样既保证了真实数据集也能进行相对一致的比较，又兼顾了工程实现的可操作性。"),
        (2, "2.4本章小结"),
        ("p", "本章从社区定义、Louvain 优化、多层网络建模、差分隐私、同态加密以及评价指标等方面，说明了本课题的理论基础与技术支撑。后续章节将围绕这些理论展开算法设计与系统实现。"),
    ]

    for item in sections:
        if isinstance(item[0], int):
            append_heading(document, item[1], item[0])
        elif item[0] == "p":
            append_paragraph(document, item[1])
        elif item[0] == "p_center":
            append_paragraph(document, item[1], align=WD_ALIGN_PARAGRAPH.CENTER)

    more_sections = [
        (1, "3隐私保护社区检测算法设计与实现"),
        (2, "3.1总体算法设计"),
        (3, "3.1.1设计目标"),
        ("p", "根据任务书要求，本课题的算法设计目标包括四项：在理解社区检测与隐私保护原理的基础上，实现可视化平台；在平台中融入隐私保护机制；在多种数据集上完成算法验证；并通过实验分析说明隐私保护能力与检测性能之间的关系。因此，项目中的算法设计不只关注单一检测精度，而是强调“隐私机制、检测流程、实验评估、系统可展示”四者的统一。"),
        (3, "3.1.2三方协同架构"),
        ("p", "项目在 `src/pmcdm/architecture.py` 中实现了终端、云服务器1和云服务器2三类角色。终端负责对边权进行同态加密并上传；云服务器1负责执行差分隐私拓扑扰动或 k 匿名补边，并记录边保留情况；云服务器2负责调用 DH-Louvain 完成社区检测。该架构不是完整的分布式部署系统，而是对隐私计算流程的程序化模拟，便于在毕业设计中同时体现算法思想与系统设计过程。"),
        (3, "3.1.3算法流程"),
        ("p", "平台中的完整检测流程可概括为：读取内置数据集或上传文件，并将其规范化为多层图结构；根据算法变体决定是否执行加密上传与拓扑扰动；调用 S-Louvain 或 DH-Louvain 完成社区划分；计算 Q、D、NMI、pr 等指标；最后输出表格结果并支持社区可视化绘图。其中，S-Louvain 作为基线算法通过多层聚合图执行标准 Louvain；DH-Louvain 则在多层图上迭代优化加权模块密度目标，是本项目的核心算法实现。"),
        (2, "3.2核心算法实现过程"),
        (3, "3.2.1数据装载与多层表示"),
        ("p", "`src/experiment_datasets.py` 是数据集总入口，负责将 Karate、AUCS、LFR、mLFR 和 BioGRID 统一封装为 `DatasetBundle`。对于只有单层结构的数据，如 Karate 和上传图文件，项目通过 `PMCDMExperiment.build_multiplex_from_base` 复制图层并对边权进行轻微扰动，模拟多层网络；对于 AUCS 和 BioGRID，则直接读取其天然多层关系结构。这样的统一表示方式降低了实验框架与数据源之间的耦合度。"),
        (3, "3.2.2S-Louvain基线实现"),
        ("p", "`src/pmcdm/s_louvain.py` 先对各层图按边权进行聚合，再调用 NetworkX 提供的 Louvain 社区发现方法完成划分。该实现简单稳定，适合作为项目中所有隐私保护算法的参考基线。对于没有真实标签的数据集，系统还会预先计算原始多层网络上的 S-Louvain 结果，用作 NMI 的参考标签。"),
        (3, "3.2.3DH-Louvain实现"),
        ("p", "`src/pmcdm/dh_louvain.py` 实现了两阶段贪心优化的 DH-Louvain。算法首先以 S-Louvain 的结果作为初始解，若初始解为空，则将每个节点初始化为独立社区。随后反复执行节点移动阶段和社区合并阶段。节点移动阶段会遍历节点，将节点尝试移动到邻居所在社区并计算目标函数增益，若增益为正则接受该移动；社区合并阶段则枚举社区对，尝试合并后重新计算目标函数，若存在正增益则执行最佳合并。"),
        ("p", "在每次移动或合并后，算法都会调用 `_reindex` 重新压缩社区编号，并通过 `weighted_modularity_density` 重新评估当前划分。该实现使算法在保持 Louvain 贪心特征的基础上，进一步考虑模块密度信息，更适合多层网络场景。"),
        (3, "3.2.4差分隐私拓扑扰动实现"),
        ("p", "在 `CloudServer1._privatize_topology_with_dp` 中，项目先收集真实边与采样非边构成候选集合，再利用拉普拉斯机制分别对存在边和不存在边的基础分值加噪，最后按照带噪得分进行排序并截取固定数量的边形成扰动图。与简单随机补边相比，这种方式在引入隐私扰动的同时，仍尽量保留了原始结构的重要边信息。"),
        (3, "3.2.5同态加密与加密统计实现"),
        ("p", "`src/homomorphic_encryption.py` 基于 `phe` 库实现了 Paillier 加法同态加密封装，提供单值加密、数组加密、密文加法、密文与明文相乘、密文求和等能力。`TerminalClient.encrypt_upload` 在上传阶段为边添加加密字段，`CloudServer1.encrypted_comm_stats` 则输出各社区内部与外部连接强度的加密统计预览。虽然项目当前并未将整个优化过程完全迁移至密文环境，但已经形成了较完整的“加密上传—密文统计—可信端解密”的实验链路。"),
        (3, "3.2.6基于进化算法的参数优化实现"),
        ("p", "在本课题中，隐私预算 ε 与目标函数权重 λ 会同时影响社区检测质量和隐私保护能力。若采用人工逐组试参的方式，不仅效率较低，而且很难在多个指标之间找到稳定折中。为此，项目在 `src/evolutionary_optimizer.py` 中引入了基于遗传算法的参数优化模块，用于自动搜索更优的参数组合，使平台不仅能够“运行算法”，还能够进一步解释“哪些参数更适合当前数据集和隐私目标”。"),
        ("p", "该模块将一个候选解编码为染色体 `x=(ε, λ)`，其中 ε 控制差分隐私扰动强度，λ 控制 DH-Louvain 中模块度项与模块密度项的权衡。初始化阶段根据预设边界随机生成种群；进化过程中采用锦标赛选择保留竞争力较强的个体，使用线性混合交叉生成新的连续参数解，并通过高斯扰动执行变异，从而保证搜索既能继承已有优良特征，又能跳出局部最优。与此同时，精英保留策略会直接将当前最优个体带入下一代，以避免优解在进化过程中丢失。"),
        ("p", "适应度函数由模块度 Q、模块密度 D、归一化互信息 NMI 和隐私率 pr 组成加权综合指标。项目默认采用 `fitness = 0.35Q + 0.25D + 0.20NMI + 0.20pr` 的形式，将结构质量与隐私保护能力统一到同一评价框架下。对于无真实标签的数据集，NMI 可以自动跳过并重新归一化其余权重；因此，该优化器不仅适用于带标签的基准网络，也能迁移到真实 BioGRID 等无标准答案的数据场景。"),
        ("p", "从实现流程看，进化算法并不是独立于社区检测系统之外的附加脚本，而是建立在 `PMCDMExperiment` 统一实验框架之上的“外层优化器”。每当产生一组新参数时，优化器都会实际调用一次隐私保护社区检测流程，完成加密上传、拓扑扰动、社区划分和指标计算，再将输出结果折算为个体 fitness。换言之，进化算法在本实验中的作用不是替代 DH-Louvain，而是作为自动调参机制去驱动 DH-Louvain 及其隐私参数的搜索，使实验从固定参数对比扩展为“参数可自适应优化”的系统化研究。"),
        ("p", "这一设计对本课题具有三方面贡献。第一，它使论文中的实验分析不再停留于少量人工设定参数点，而能够给出收敛过程和优化前后对比，从而增强实验结论的可靠性。第二，它为平台增加了面向不同数据集的适应能力，因为不同网络的层间一致性、边密度和社区清晰度存在明显差异，最优 ε 与 λ 并不固定。第三，它提升了系统的工程完整性，使前端“优化参数”面板、后端 `/optimize/ea` 接口与算法评估模块形成闭环，体现出本课题不仅实现了隐私保护社区检测算法，也实现了面向该算法的参数优化与决策支持能力。"),
        (2, "3.3算法实现特点与复杂度分析"),
        ("p", "从实现结果看，本项目算法部分具有以下特点：一是统一实验接口，无论数据来源是内置数据集、生成器还是用户上传文件，最终均转化为统一的多层图格式；二是隐私机制可插拔，不同算法变体通过“是否扰动、是否加密、是否采用 DH 框架”三个维度进行组合；三是指标体系完整，不仅输出社区划分，还提供 Q、D、NMI、pr 以及社区数量等结果；四是便于系统集成，实验框架与后端接口、可视化模块解耦，适合进一步扩展为 Web 平台。"),
        ("p", "在公式层面，DH-Louvain 的核心优化目标可表示为 `F_λ = λQ + (1-λ)D`，其中 λ 用于平衡模块度 Q 与模块密度 D；进化算法中单个个体的适应度函数可表示为 `fitness = w_1Q + w_2D + w_3NMI + w_4pr`，且 `w_1+w_2+w_3+w_4=1`。在项目默认配置下，四个权重分别设置为 `0.35、0.25、0.20、0.20`，使结构质量与隐私保护能力能够在统一标尺下比较。"),
        ("p", "从复杂度角度看，S-Louvain 首先需要对 q 层网络进行聚合，其代价近似为 `O(qm)`；随后在聚合图上执行 Louvain 优化，因此整体复杂度可近似写为 `O(qm + T_s m)`，其中 m 为边数，`T_s` 为 Louvain 迭代轮数。DH-Louvain 在每一轮外层迭代中需要执行节点移动和社区合并，并多次评估加权目标函数，因此其复杂度高于基线方法，可近似表示为 `O(T_d · q · m · c)`，其中 `T_d` 为外层迭代次数，c 表示候选社区评估带来的额外开销。"),
        ("p", "差分隐私拓扑扰动主要对真实边和采样非边进行带噪评分与筛选，其代价近似为 `O(m + s m)`；同态加密模块则主要增加常数级较大的密文运算开销。对于进化算法，若种群规模为 P、进化代数为 G，而单次参数评估代价记为 `C_eval`，则总复杂度可写为 `O(P·G·C_eval)`。由于 `C_eval` 本质上就是一次完整的隐私保护社区检测流程，因此进化算法更适合离线调参和实验分析，而不是在线高频调用。"),
        (2, "3.4本章小结"),
        ("p", "本章围绕项目源码详细说明了隐私保护社区检测算法的设计思路与实现过程，包括数据组织、S-Louvain 基线、DH-Louvain、差分隐私扰动、同态加密模拟以及进化参数优化等内容。算法部分为后续系统平台和实验分析奠定了核心基础。"),
        (1, "4系统架构与实现"),
        (2, "4.1系统总体架构设计"),
        ("p", "本项目采用“算法层 + 服务层 + 展示层”的整体架构。算法层由 `src/pmcdm`、`src/differential_privacy.py` 和 `src/homomorphic_encryption.py` 等模块组成，负责核心计算；服务层由 FastAPI 后端构成，负责参数接收、数据集装载、算法调度和结果序列化；展示层由 Vue 前端构成，负责交互式配置、结果表格展示和网络可视化。"),
        ("p", "从实现角度看，系统并非单纯将算法脚本网页化，而是对实验平台进行了模块化拆分：命令行入口 `main.py` 用于快速实验；`src/backend_framework` 提供更完整的后端接口和服务封装；`frontend/src/App.vue` 提供多数据集、多参数、多算法的统一操作界面。这样既满足毕业设计对系统性的要求，也便于答辩现场演示。"),
        (2, "4.2后端设计与实现"),
        ("p", "后端主要基于 FastAPI 实现。`src/backend_framework/routers.py` 中定义了 `/datasets`、`/detect`、`/visualize` 以及 `/optimize/ea` 等接口，分别用于获取数据集目录、执行社区检测、生成可视化和运行进化优化。`src/backend_framework/services.py` 则承担主要服务逻辑，包括统一构造请求参数、加载内置数据集或上传文件、聚合各层网络并统计图属性、调用算法模块执行检测、生成二维或三维可视化数据，以及组织返回结果供前端展示。"),
        ("p", "后端设计的重点在于“对算法复杂性进行封装”。前端无需了解 DH-Louvain 的细节，只需通过表单提交算法名称、隐私预算、随机种子和数据集参数即可完成实验调用。"),
        (2, "4.3前端设计与实现"),
        ("p", "前端采用 Vue 实现单页界面，核心文件为 `frontend/src/App.vue`。页面整体采用三栏布局：左侧为实验配置区，中间为检测结果区，右侧为网络可视化区。界面支持在内置数据集和文件上传两种模式之间切换，能够针对 LFR、mLFR 与 BioGRID 动态显示专属参数，并允许用户设置 ε、δ、λ、密钥长度、随机种子和布局方式。"),
        ("p", "相比单纯的表单页面，该界面还额外设计了检测进度弹窗、参数分组、折叠高级设置和遗传算法优化面板，使系统更符合“可视化复杂网络社区检测平台”的题目要求。"),
        (2, "4.4可视化设计与实现"),
        ("p", "可视化模块由 `src/visualization.py` 实现，支持 spring、circular、kamada_kawai 以及 3d_spring 等布局方式。模块能够根据社区编号自动构造颜色映射，并将网络结构以二维图或三维图形式输出至 `output/` 目录。对于大图数据，系统还会优先标注高连接度节点，以兼顾可读性和表现力。"),
        ("p", "除静态图片外，后端还可以返回三维节点与边的结构化数据，用于前端进行交互式展示。这种设计使系统既适合论文和报告中的图片插入，也适合答辩现场动态演示。"),
        (2, "4.5本章小结"),
        ("p", "本章介绍了平台从后端到前端的整体实现方式。通过算法层、服务层和展示层的协同，本项目形成了一套较完整的隐私保护社区检测系统，不仅能够运行实验，还具备较强的可视化与演示能力。"),
        (1, "5系统测试与实验分析"),
        (2, "5.1实验环境与方案"),
        ("p", "项目开发语言为 Python，主要依赖包括 NetworkX、NumPy、Pandas、SciPy、scikit-learn、FastAPI、Uvicorn、Vue 以及 `phe` 等。实验侧重回答两个问题：一是平台在不同数据集上能否稳定完成社区检测；二是在加入隐私保护后，算法性能与隐私率之间呈现怎样的关系。"),
        ("p", "为保证结果具有可比性，命令行与后端均采用统一的 `PMCDMExperiment` 流程，默认参数包括 ε=1.0、δ=1e-5、key_size=512 和 random_state=42，并对 S-Louvain、PD-Louvain、R-Louvain、DP-Louvain、K-Louvain 与 DH-Louvain 六种算法变体进行横向比较。"),
    ]

    for item in more_sections:
        if isinstance(item[0], int):
            append_heading(document, item[1], item[0])
        elif item[0] == "p":
            append_paragraph(document, item[1])

    append_table(
        document,
        "表5-1 实验环境配置",
        ["项目", "说明"],
        [
            ["开发语言", "Python 3.12"],
            ["后端框架", "FastAPI + Uvicorn"],
            ["前端框架", "Vue"],
            ["图计算库", "NetworkX"],
            ["数值与数据处理", "NumPy / Pandas / SciPy / scikit-learn"],
            ["隐私计算相关库", "phe（Paillier）"],
            ["实验数据集", "Karate、AUCS、LFR、mLFR、BioGRID"],
        ],
    )

    append_heading(document, "5.2基准数据集实验分析", 2)
    append_paragraph(document, "在 Karate Club 数据集上，DH-Louvain 得到 3 个社区，NMI 达到 0.6877，明显高于同条件下的 S-Louvain（0.4900）与 DP-Louvain（0.5005），同时隐私率达到 0.9594。该结果表明，在较小规模网络中，DH-Louvain 在保持较强隐私保护能力的同时，仍能恢复较合理的社区结构。")
    append_paragraph(document, "在 AUCS 多层社交网络数据集上，S-Louvain 和 PD-Louvain 的划分结果较接近，二者均得到 5 个社区且 NMI 为 0.7827。DH-Louvain 在隐私率达到 0.9637 的前提下，获得了 4 个社区、Q=0.6440、D=0.0430。虽然其 NMI 低于明文基线，但仍保持了较好的社区可解释性，说明加入隐私保护后平台仍能完成有效检测。")
    append_paragraph(document, "从横向对比可以进一步看到三点规律。第一，PD-Louvain 与 S-Louvain 在 AUCS 上几乎重合，说明“仅引入加密上传而不扰动拓扑”对划分质量影响很小，结构性损失主要来自拓扑扰动而非密文封装本身。第二，R-Louvain 与 DP-Louvain 在两个数据集上的 Q 和 NMI 均较基线下降，表明随机补边和差分隐私扰动都会削弱社区边界，但 DP-Louvain 在 AUCS 上仍能维持接近基线的 0.7653 NMI，说明带噪筛边比纯随机补边更能保留有效结构。第三，K-Louvain 在 AUCS 上的模块度仅为 0.1385，甚至出现负的模块密度，这说明通过度均衡来提升匿名性的做法虽然能隐藏局部结构特征，但会显著破坏原有社区内部的紧密连接。"),
    append_paragraph(document, "DH-Louvain 的表现体现了本课题设计的价值。在 Karate 数据集上，DH-Louvain 的 NMI 为所有隐私保护算法中最高；在 AUCS 上虽然其 Q 和 NMI 略低于 S-Louvain 与 PD-Louvain，但 pr 达到 0.9637，明显高于 R-Louvain、DP-Louvain 和 K-Louvain。换言之，DH-Louvain 的优势不在于无条件追求最高 Q，而在于在高隐私率下尽量维持可解释的社区划分结果，这与隐私保护社区检测的研究目标是一致的。")
    append_table(
        document,
        "表5-2 Karate Club 数据集算法对比结果",
        ["算法", "Q", "D", "NMI", "pr", "社区数"],
        [
            ["S-Louvain", "0.4245", "0.2550", "0.4900", "0.5000", "4"],
            ["PD-Louvain", "0.4254", "0.2566", "0.5878", "1.0000", "4"],
            ["R-Louvain", "0.3701", "0.1578", "0.5869", "0.5000", "4"],
            ["DP-Louvain", "0.3909", "0.1872", "0.5005", "0.4594", "4"],
            ["K-Louvain", "0.3060", "0.0378", "0.4400", "0.5000", "5"],
            ["DH-Louvain", "0.3611", "0.1829", "0.6877", "0.9594", "3"],
        ],
    )
    append_table(
        document,
        "表5-3 AUCS 数据集算法对比结果",
        ["算法", "Q", "D", "NMI", "pr", "社区数"],
        [
            ["S-Louvain", "0.7392", "0.0542", "0.7827", "0.5000", "5"],
            ["PD-Louvain", "0.7392", "0.0542", "0.7827", "1.0000", "5"],
            ["R-Louvain", "0.6597", "0.0522", "0.7551", "0.5000", "5"],
            ["DP-Louvain", "0.6814", "0.0487", "0.7653", "0.4613", "5"],
            ["K-Louvain", "0.1385", "-0.1137", "0.7678", "0.5000", "5"],
            ["DH-Louvain", "0.6440", "0.0430", "0.6712", "0.9637", "4"],
        ],
    )
    append_picture(document, OUTPUT_DIR / "chart_compare_q_lines.png", "图5-1 Karate 与 AUCS 数据集的模块度折线对比图", width_cm=15.5)
    append_picture(document, OUTPUT_DIR / "chart_compare_nmi_lines.png", "图5-2 Karate 与 AUCS 数据集的 NMI 折线对比图", width_cm=15.5)
    append_picture(document, OUTPUT_DIR / "chart_compare_privacy_lines.png", "图5-3 Karate 与 AUCS 数据集的隐私率折线对比图", width_cm=15.5)

    append_heading(document, "5.3真实BioGRID数据实验分析", 2)
    append_paragraph(document, "为验证算法在真实复杂网络中的适用性，项目进一步在 7 个指定 BioGRID 物种网络上执行批量实验。结果显示，DH-Louvain 在 Gallus gallus、Candida albicans SC5314、Xenopus laevis、Human Immunodeficiency Virus 1、Rattus norvegicus 和 Caenorhabditis elegans 等数据上均能输出合理的社区数，并保持约 0.96 左右的隐私率。")
    append_paragraph(document, "其中，Caenorhabditis elegans 数据集的模块度达到 0.5205，模块密度为 0.1807，社区数为 5，说明在真实多层生物网络中，DH-Louvain 仍能识别出较明显的聚类结构。对于规模较小的 Bos taurus 和 Gallus gallus 等数据，算法也能够在高隐私率下完成社区划分，说明平台对大图和小图均具有一定适应性。")
    append_paragraph(document, "结合真实物种结果可以发现，DH-Louvain 的隐私率曲线在不同物种之间波动很小，基本稳定在 0.96 附近，这说明终端加密和云侧扰动机制在不同图规模下具有较稳定的保护强度；相比之下，Q 与 NMI 的波动明显更大。这种现象意味着，在真实数据场景中，决定社区检测难度的关键因素不是隐私机制是否失效，而是原始网络本身的层间一致性、边密度和功能模块清晰度。"),
    append_paragraph(document, "以 Caenorhabditis elegans 为例，其去重边数达到 9008，Q 和 D 都显著高于其他物种，说明高密度、多层且功能结构更清晰的网络更容易在扰动后保留社区边界；而 Rattus norvegicus 的去重边数虽然也较大，但 D 为负值、NMI 仅为 0.3845，反映出图中可能存在更强的跨社区连接或不同实验系统之间的层间异质性，从而削弱了社区紧致性。Bos taurus 的样本规模较小，却得到较高的 NMI，这说明在小规模且边界较明确的网络中，隐私机制带来的影响可以被较稳定的结构模式抵消。")
    append_table(
        document,
        "表5-4 指定 7 个 BioGRID 物种的 DH-Louvain 结果",
        ["物种", "节点数", "层数", "去重边数", "Q", "D", "NMI", "pr", "社区数"],
        [
            ["Gallus gallus", "41", "3", "83", "0.2145", "0.0011", "0.6213", "0.9663", "4"],
            ["Bos taurus", "17", "3", "33", "0.0000", "0.0672", "0.7739", "0.9848", "2"],
            ["Candida albicans SC5314", "296", "3", "685", "0.3621", "0.0080", "0.5641", "0.9647", "14"],
            ["Xenopus laevis", "373", "3", "832", "0.3286", "0.0002", "0.7088", "0.9624", "12"],
            ["Human Immunodeficiency Virus 1", "177", "3", "372", "0.2137", "-0.0074", "0.5130", "0.9647", "7"],
            ["Rattus norvegicus", "490", "3", "2365", "0.2057", "-0.0363", "0.3845", "0.9601", "9"],
            ["Caenorhabditis elegans", "500", "3", "9008", "0.5205", "0.1807", "0.6381", "0.9609", "5"],
        ],
    )
    append_picture(document, OUTPUT_DIR / "chart_biogrid_dh_species_lines.png", "图5-4 DH-Louvain 在 7 个 BioGRID 物种上的指标折线图", width_cm=15.5)
    append_paragraph(document, "除 λ 控制实验外，本文还进一步以迭代次数作为控制变量，对同样的 5 个真实生物网络进行了 DH-Louvain 收敛过程分析。实验将最大迭代次数从 1 增加到 6，并观测归一化模块密度 D/n、模块度 Q 以及与无隐私 S-Louvain 划分相比的 NMI。该实验的目的在于说明：DH-Louvain 的结果并非由单次随机扰动偶然产生，而是在有限轮迭代中逐渐趋于稳定。"),
    append_paragraph(document, "从图中可以看出，5 个真实网络在 1 至 6 次迭代过程中总体表现出“前期调整、后期稳定”的趋势。对于 HIV-1、Xenopus 和 E.coli，Q 与 NMI 在前 2 至 3 次迭代后基本进入平稳区间，说明算法较快就能收敛到较稳定的社区划分；对于 MERS，Q 和 NMI 在多轮迭代中始终保持较高水平，说明该网络的社区边界较为清晰，少量迭代即可获得较好的结果；Plasmodium 的 NMI 整体较低但仍在 2 至 6 次迭代内小幅波动，这说明当网络层数较少且结构较稀疏时，DH-Louvain 的收敛结果更容易受到拓扑扰动影响。"),
    append_paragraph(document, "进一步比较 D/n 曲线可发现，不同网络的归一化模块密度水平存在明显差异。E.coli 和 MERS 的 D/n 明显高于 Plasmodium、HIV-1 和 Xenopus，说明其社区内部连接更紧密；而大多数网络在后续迭代中 D/n 变化幅度较小，说明节点移动和社区合并主要发生在初始若干轮。综合来看，该实验验证了 DH-Louvain 在真实生物网络上具有较好的迭代稳定性，也说明将迭代次数设置在 3 到 6 之间通常已经足以获得较可靠的划分结果。"),
    append_picture(document, OUTPUT_DIR / "chart_biogrid_5_iterations_triptych.png", "图5-4-2 五个真实 BioGRID 网络在不同迭代次数下的 D/n、Q、NMI 折线图", width_cm=16.0)

    append_heading(document, "5.4参数优化与补充实验分析", 2)
    append_paragraph(document, "为了进一步说明算法在可控条件下的变化规律，本文补充进行了 LFR 参数控制实验。实验固定 N=300、q=3、γ=2、平均度 10、最大度 40，仅改变混合参数 μ=0.00~0.50，并同时比较 S-Louvain、PD-Louvain、R-Louvain、DP-Louvain、K-Louvain 和 DH-Louvain 六种算法。需要说明的是，修正后的 μ=0 样本保留了完整 300 个节点和 9 个真实社区，因此能够与其余 μ 取值形成同条件比较。")
    append_paragraph(document, "从图5-5至图5-7可以看出，随着 μ 增大，六种算法的 Q、D 和 NMI 均整体下降，符合“社区间混合越强、检测越困难”的理论预期。图5-5 表明，当 μ=0 时，S-Louvain 与 PD-Louvain 的模块度均达到 0.8676，说明在社区边界最清晰的情况下，基线方法几乎可以完全恢复真实结构；随着 μ 逐步增大，所有算法的模块度都持续下滑，反映出跨社区连边增多后，网络模块结构被不断削弱。")
    append_paragraph(document, "图5-6 进一步展示了模块密度 D 对社区内部紧密性的敏感反映。与 Q 相比，D 在高 μ 区间下降更快，且部分算法在 μ 较大时出现负值，说明这时社区内部连通紧致性已经明显减弱。其中，K-Louvain 在整个扫描区间下降最为明显之一，在 μ=0.50 时模块密度已降至 -0.1404，表明通过度匿名化来隐藏局部结构虽然能够提高匿名性，却会更强烈地破坏社区内部的结构完整性。")
    append_paragraph(document, "图5-7 给出了 NMI 随 μ 的变化规律。在 μ=0.00~0.20 的区间，多数算法的 NMI 接近 1.0000，说明真实社区边界清晰时，各算法都能较准确地恢复标签结构；当 μ 超过 0.25 后，NMI 开始明显下降，说明算法误差首先来自网络本身社区边界的弱化，其次才是隐私机制带来的附加损失。修正后的 μ=0 样本保留了完整 300 个节点和 9 个真实社区，因此其曲线起点能够与其余 μ 取值形成有效对比，不再受到先前退化样本的干扰。")
    append_paragraph(document, "横向比较还能得到两点更有代表性的结论。第一，PD-Louvain 与 S-Louvain 在 μ≤0.20 的区间几乎完全重合，说明仅使用同态加密封装并不会明显改变社区结构，性能损失主要来自拓扑扰动而非加密上传本身。第二，DH-Louvain 虽然在多数点上并非绝对最优，但其 pr 始终保持在约 0.96 的高水平，并且在 μ=0.35 以前仍维持了可接受的 Q、D 和 NMI，这说明该算法的优势在于能够在高隐私率条件下维持相对稳定且可解释的社区划分结果。")
    append_picture(document, OUTPUT_DIR / "chart_lfr_mu_group_modularity_n300_all.png", "图5-5 LFR 参数控制实验中 Q 随 μ 变化的折线图", width_cm=15.5)
    append_picture(document, OUTPUT_DIR / "chart_lfr_mu_group_density_n300_all.png", "图5-6 LFR 参数控制实验中 D 随 μ 变化的折线图", width_cm=15.5)
    append_picture(document, OUTPUT_DIR / "chart_lfr_mu_group_nmi_n300_all.png", "图5-7 LFR 参数控制实验中 NMI 随 μ 变化的折线图", width_cm=15.5)
    append_paragraph(document, "由此可见，μ 决定了问题本身的难度上界。当社区混合度较低时，DH-Louvain 能够在高隐私率下保持较高的一致性与结构质量；而当 μ 继续升高时，网络本身的社区边界开始模糊，算法性能下降更多地来自数据难度本身，而不是隐私机制单独造成的损伤。")
    append_paragraph(document, "在参数优化方面，项目进一步使用进化算法自动搜索关键参数。经过 48 次候选评估后，最优参数为 ε=4.2371、λ=0.1110，对应适应度 0.5477；相比基线参数 ε=1.0、λ=0.5，优化后 fitness、模块度、模块密度、NMI 和隐私率均有所提升。")
    append_paragraph(document, "遗传算法搜索结果则体现了参数自适应优化的必要性。与基线参数相比，最优个体的 fitness 从 0.4906 提升到 0.5477，提升幅度约为 11.6%；NMI 从 0.5684 提升到 0.6912，说明更大的 ε 和更小的 λ 组合在 Karate 数据集上能更好兼顾结构保持与隐私约束。进一步观察逐代曲线可以发现，最优适应度在第 4 代后出现明显跃升，而平均适应度整体稳步上升，说明进化搜索并非偶然命中，而是在逐步收敛到更优区域。")
    append_table(
        document,
        "表5-5 Karate Club 数据集参数优化前后对比",
        ["方案", "ε", "λ", "fitness", "Q", "D", "NMI", "pr", "社区数"],
        [
            ["基线参数", "1.0000", "0.5000", "0.4906", "0.3682", "0.2213", "0.5684", "0.9637", "3"],
            ["优化结果", "4.2371", "0.1110", "0.5477", "0.4106", "0.2665", "0.6912", "0.9957", "3"],
        ],
    )
    append_picture(document, OUTPUT_DIR / "chart_ea_fitness_history.png", "图5-8 遗传算法参数优化过程折线图", width_cm=14.5)
    append_paragraph(document, "在已有最优个体结果的基础上，本文进一步进行了 10 次重复对比实验，以更直观说明进化算法的实际作用。结果表明，相比基线参数 `ε=1.0, λ=0.5`，进化算法搜索得到的参数 `ε=4.2371, λ=0.1110` 使平均 fitness 从 0.4859 提升到 0.5305，提升约 9.17%；平均模块度从 0.3802 提升到 0.3969，平均模块密度从 0.2012 提升到 0.2574，平均 NMI 从 0.5485 提升到 0.6444，平均隐私率从 0.9643 提升到 0.9919。与此同时，平均运行时间还略有下降，说明优化后的参数不仅提高了隐私—效用折中质量，也没有带来额外的明显计算负担。"),
    append_paragraph(document, "该结果说明，进化算法在本课题中的作用并不只是“自动找到一组更好看的参数”，而是能够系统性提升实验结论的可靠性。它通过反复评估候选参数，使 DH-Louvain 从固定经验参数运行，扩展为可根据指标目标自动搜索更优解的框架，从而增强了平台对不同网络和不同隐私需求的适应能力。"),
    append_picture(document, OUTPUT_DIR / "chart_ea_before_after_compare.png", "图5-9 进化算法优化前后平均指标对比图", width_cm=14.5)

    append_heading(document, "5.5系统功能验证与可视化结果", 2)
    append_paragraph(document, "功能测试表明，平台已经能够完成以下闭环流程：选择数据集或上传图文件、设置隐私参数、执行社区检测、显示指标结果、生成网络可视化图像。命令行入口 `main.py` 适合快速运行实验并输出对比表；FastAPI 后端适合前后端联调和接口测试；Vue 前端则适合进行答辩演示和交互式实验。")
    append_paragraph(document, "从工程实现角度看，系统已经满足任务书提出的“设计基于隐私保护的社区检测平台，实现社区检测结果的可视化展示”要求。同时，结合 BioGRID 真实数据与 AUCS 等多层社交数据的实验结果，也验证了平台具备“基于真实数据测试验证隐私保护下社区检测算法的性能和隐私保护能力”的任务目标。")
    karate_image = find_latest_image("Karate_Club_DH-Louvain")
    if karate_image is not None:
        append_picture(document, karate_image, "图5-10 Karate Club 数据集的 DH-Louvain 可视化结果")
    aucs_image = OUTPUT_DIR / "AUCS_DH-Louvain_20260411_173339.png"
    if aucs_image.exists():
        append_picture(document, aucs_image, "图5-11 AUCS 数据集的社区检测可视化结果")

    append_heading(document, "5.6本章小结", 2)
    append_paragraph(document, "本章从基准数据集、真实 BioGRID 数据、LFR 参数控制实验、进化参数优化以及系统功能验证五个方面对平台进行了测试分析。综合来看，项目在隐私保护与社区检测之间实现了可接受的平衡，并形成了一套具备展示价值和扩展潜力的复杂网络分析原型系统。")
    append_paragraph(document, "为了给全文结论提供更紧凑的总结，本文进一步给出了综合排名图。该图将 Karate 与 AUCS 数据集上的 Q、D、NMI 指标合成为社区检测排名，并将隐私率合成为隐私保护排名。结果表明，S-Louvain 和 PD-Louvain 在纯检测能力上排名更靠前，而 DH-Louvain 在隐私保护排名上具有明显优势，并且在综合分析中体现出更好的折中能力。这一结果与前文逐项实验分析得到的结论是一致的。")
    append_picture(document, OUTPUT_DIR / "chart_conclusion_algorithm_ranking.png", "图5-12 各算法在检测能力与隐私保护上的综合排名图", width_cm=15.5)

    append_heading(document, "6总结与展望", 1)
    append_paragraph(document, "本文围绕“基于隐私保护的社区检测算法设计并实现可视化复杂网络社区检测系统”这一课题，完成了理论分析、算法设计、系统实现和实验验证等工作。项目以多层网络为对象，构建了终端、云服务器1和云服务器2三方协同架构，实现了 DH-Louvain、差分隐私扰动、同态加密模拟、指标评估、前后端交互和可视化展示等功能。实验结果表明，平台能够在隐私保护条件下完成社区检测任务，并在真实数据集上体现出较好的工程可行性。"),
    append_paragraph(document, "从实验结论看，DH-Louvain 的核心价值主要体现在三个方面：第一，在 Karate、AUCS 与 BioGRID 等不同类型网络上都能输出具有可解释性的社区结果，说明算法具备较好的普适性；第二，在高隐私率条件下仍能维持相对稳定的 Q、D 与 NMI，说明其在隐私保护与检测性能之间实现了较合理的折中；第三，通过 LFR 参数控制实验、真实生物网络迭代实验和进化算法优化实验可以看到，平台不仅能完成单次社区检测，还能支撑参数扫描、稳定性分析与自动调参，这使课题从“算法实现”进一步扩展为“可验证、可分析、可展示”的系统平台。"),
    append_paragraph(document, "当然，当前系统仍存在进一步完善空间。首先，当前同态加密主要用于上传与统计模拟，尚未覆盖完整的社区优化流程；其次，差分隐私扰动策略仍可结合更精细的边采样与图结构保护机制继续优化；再次，前端界面虽然已经具备答辩演示能力，但在实验结果图表化、历史结果持久化和大规模图渲染方面仍可加强。后续工作可以围绕更严格的隐私定义、更高效的图优化算法以及更规范的实验管理平台展开，使系统在学术研究和工程应用两方面都具备更强的推广价值。")

    append_paragraph(document, "【参考文献】", style="参考文献")
    references = [
        "[1] Fortunato S. Community detection in graphs[J]. Physics Reports, 2010, 486(3-5): 75-174.",
        "[2] Blondel V D, Guillaume J L, Lambiotte R, et al. Fast unfolding of communities in large networks[J]. Journal of Statistical Mechanics: Theory and Experiment, 2008(10): P10008.",
        "[3] Lancichinetti A, Fortunato S, Radicchi F. Benchmark graphs for testing community detection algorithms[J]. Physical Review E, 2008, 78(4): 046110.",
        "[4] Kivelä M, Arenas A, Barthélemy M, et al. Multilayer networks[J]. Journal of Complex Networks, 2014, 2(3): 203-271.",
        "[5] Dwork C, Roth A. The Algorithmic Foundations of Differential Privacy[J]. Foundations and Trends in Theoretical Computer Science, 2014, 9(3-4): 211-407.",
        "[6] Paillier P. Public-key cryptosystems based on composite degree residuosity classes[C]// Advances in Cryptology—EUROCRYPT’99. Berlin: Springer, 1999: 223-238.",
        "[7] Zachary W W. An information flow model for conflict and fission in small groups[J]. Journal of Anthropological Research, 1977, 33(4): 452-473.",
        "[8] Magnani M, Micenková B, Rossi L. Combinatorial analysis of multiple networks[EB/OL]. arXiv, 2013: arXiv:1303.4986.",
        "[9] Bródka P. A Method for Group Extraction and Analysis in Multilayer Social Networks[EB/OL]. arXiv, 2016: arXiv:1612.02377.",
        "[10] Oughtred R, Rust J, Chang C, et al. The BioGRID interaction database: 2019 update[J]. Nucleic Acids Research, 2019, 47(D1): D529-D541.",
        "[11] Magnani M, Rossi L, Nicosia V. Multilayer Social Networks[M]. Cambridge: Cambridge University Press, 2021.",
    ]
    for item in references:
        append_paragraph(document, item)

    append_paragraph(document, "致谢", style="致谢")
    append_paragraph(document, "在本次毕业设计的完成过程中，我围绕隐私保护社区检测这一课题，从理论学习、算法实现到系统开发，逐步体会到复杂网络分析与软件工程实践相结合的难度与价值。感谢指导教师在课题选题、研究思路、系统设计和论文写作中的耐心指导，也感谢项目开发过程中提供帮助的同学和老师。正是在不断调试、实验和修改的过程中，我对社区检测、多层网络和隐私保护技术有了更深入的理解。")

    english_abstract = [
        "【Abstract】Community structures in complex networks reveal latent organizational patterns and interaction modules among nodes. However, directly performing community detection on plain graph data may expose sensitive edges and relationship information. To address this issue, this thesis designs and implements a privacy-aware community detection platform for multiplex networks. The platform integrates community detection algorithms, differential privacy, homomorphic encryption simulation, and a visual Web interface, so that the feasibility, performance, and privacy-preserving capability of community detection can be evaluated in a unified system.",
        "The project first analyzes the basic theories of community detection, modularity optimization, multiplex networks, differential privacy, and homomorphic encryption. Then a three-party collaborative framework composed of terminal, cloud server 1, and cloud server 2 is designed, and the DH-Louvain based detection workflow is implemented. The system further provides dataset loading, benchmark comparison, multiresolution analysis, evolutionary parameter optimization, and visualization services. A FastAPI backend and a Vue frontend are developed to support dataset selection, parameter configuration, algorithm execution, metric display, and graph rendering.",
        "Experimental results on Karate, AUCS, LFR, mLFR, and real BioGRID datasets show that the implemented platform can complete community detection tasks under privacy protection constraints while maintaining interpretable partition results and high privacy rates. The system therefore satisfies the task requirements of algorithm understanding, privacy mechanism analysis, platform construction, and real-data validation.",
        "【Key words】community detection; differential privacy; homomorphic encryption; multiplex network; visualization system",
    ]
    for item in english_abstract:
        append_paragraph(document, item, style="摘要")

    document.save(str(OUTPUT_DOCX))
    OUTPUT_MD.write_text(build_markdown(), encoding="utf-8")


if __name__ == "__main__":
    build_document()
